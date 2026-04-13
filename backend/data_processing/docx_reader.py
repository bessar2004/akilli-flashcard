"""
data_processing/docx_reader.py
--------------------------------
python-docx kullanarak DOCX dosyasından ham metin çıkarır.
"""

from docx import Document as DocxDocument


def extract_text_from_docx(file_path: str) -> str:
    """
    Verilen DOCX dosyasının tüm paragraflarından metin çıkarır.

    Args:
        file_path: DOCX dosyasının mutlak yolu.

    Returns:
        Paragrafların birleştirilmiş ham metni.

    Raises:
        RuntimeError: Dosya açılamazsa.
    """
    try:
        doc = DocxDocument(file_path)
    except Exception as e:
        raise RuntimeError(f"DOCX açılamadı: {e}")

    paragraphs: list[str] = []

    # Gövde paragrafları
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraphs.append(text)

    # Tablolardaki metinleri de dahil et
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in paragraphs:
                    paragraphs.append(text)

    if not paragraphs:
        raise RuntimeError("DOCX'ten metin çıkarılamadı. Dosya boş olabilir.")

    return "\n\n".join(paragraphs)
